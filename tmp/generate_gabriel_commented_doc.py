from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/pablogalindo/Desktop/UNIVERSIDAD/NOVENO SEMESTRE/Móviles/G36_Moviles")
OUTPUT_DIR = ROOT / "deliverables"
OUTPUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUTPUT_DIR / "Avances_Gabriel_y_pendientes_comentado.docx"


GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xC6, 0x28, 0x28)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK = RGBColor(0x22, 0x22, 0x22)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_FILL = "E8EEF5"


def set_run_font(run, size, bold=False, color=DARK, name="Calibri", italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_para_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, inches):
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(inches * 1440)))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def build_styles(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78), 10, 5),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    if "Comment Body" not in document.styles:
        style = document.styles.add_style("Comment Body", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = document.styles["Comment Body"]
    style.base_style = document.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(9)
    style.font.color.rgb = GRAY
    style.paragraph_format.left_indent = Inches(0.32)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15


def add_title(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, before=0, after=4, line=1.0)
    run = p.add_run("Sprint4_Checklist")
    set_run_font(run, 24, bold=True)

    p = document.add_paragraph()
    set_para_spacing(p, before=0, after=10, line=1.1)
    run = p.add_run("Checklist Sprint 4 - Fluxo (Team G36)")
    set_run_font(run, 20, bold=True)

    for label, value in [
        ("Deadline", "Viernes 23 May 2026, 5:00 AM (GMT-5)"),
        ("Ultima actualizacion original", "20 May 2026"),
        ("Equipo", "Gabriel Padilla, Companero A, Companero B"),
        ("Version comentada", "21 May 2026 - verificacion y comentarios de implementacion iOS"),
    ]:
        p = document.add_paragraph()
        set_para_spacing(p, before=0, after=2, line=1.0)
        r1 = p.add_run(f"{label}: ")
        set_run_font(r1, 11, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, 11)


def add_item(document, text, done, comment):
    p = document.add_paragraph()
    set_para_spacing(p, before=0, after=1, line=1.15)
    symbol = "☑ " if done else "☐ "
    run = p.add_run(symbol)
    set_run_font(run, 11, bold=True, color=GREEN if done else RED)
    run = p.add_run(text)
    set_run_font(run, 11, bold=False)

    cp = document.add_paragraph(style="Comment Body")
    set_para_spacing(cp, before=0, after=4, line=1.15)
    run = cp.add_run("Comentario: ")
    set_run_font(run, 9, bold=True, color=GRAY)
    run = cp.add_run(comment)
    set_run_font(run, 9, color=GRAY)


def add_bullet(document, text, indent=0.5):
    p = document.add_paragraph()
    set_para_spacing(p, before=0, after=3, line=1.15)
    p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run("• ")
    set_run_font(run, 10, bold=False)
    run = p.add_run(text)
    set_run_font(run, 10)


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_width(hdr[idx], widths[idx])
        set_cell_margins(hdr[idx])
        shade_cell(hdr[idx], LIGHT_FILL)
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = hdr[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_spacing(paragraph, before=0, after=0, line=1.0)
        run = paragraph.add_run(header)
        set_run_font(run, 10, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_width(cells[idx], widths[idx])
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[idx].paragraphs[0]
            set_para_spacing(paragraph, before=0, after=0, line=1.1)
            run = paragraph.add_run(value)
            set_run_font(run, 9.5)
    document.add_paragraph()


def add_sections(document):
    document.add_heading("Punto 2 - Value Proposition Refinement (10 pts)", level=1)
    add_item(document, "Documentar value proposition refinada en Wiki", False, "Pendiente manual. No se puede cerrar desde el repo; requiere escritura en la wiki.")
    add_item(document, "Incluir feedback recibido en Sprint 3 si lo hubo", False, "Pendiente manual. Solo puede cerrarse con la version final de la wiki.")

    document.add_heading("Punto 3 - Micro-optimization (40 pts) - Companeros", level=1)
    add_item(document, "(Companeros) Identificar 1-2 lugares con mala UX (slow load, jank)", False, "Pendiente manual de analisis/profiling. En codigo ya se aplico una optimizacion en dashboard, pero la seleccion formal de hotspots debe documentarse.")
    add_item(document, "(Companeros) Profiling con Instruments ANTES (screenshot guardado)", False, "Pendiente manual. No queda evidencia de Instruments dentro del repo.")
    add_item(document, "(Companeros) Aplicar optimizacion (LazyVStack, async let en dashboard, decode en background, etc.)", True, "Hecho en iOS. Se aplico LazyVStack en Features/Dashboard/DashboardView.swift y carga concurrente con async let en Features/Dashboard/DashboardViewModel.swift.")
    add_item(document, "(Companeros) Profiling DESPUES (screenshot guardado)", False, "Pendiente manual. Falta la corrida posterior en Instruments y sus capturas.")
    add_item(document, "(Companeros) Documentar mejoras medidas en Wiki", False, "Pendiente manual. El cambio de codigo ya esta listo, pero falta documentar metricas y comparacion en wiki.")

    document.add_heading("Punto 4 - Three Features (80 pts total)", level=1)
    document.add_heading("F1 - CRUD Expenses", level=2)
    add_item(document, "1.2 Backend - PostgREST + RLS para update/delete", True, "Verificado en el repo. Ya estaba implementado antes y no fue modificado en esta iteracion.")
    add_item(document, "1.3 ExpensesAdapter - updateExpense + deleteExpense + ExpenseUpdatePatch", True, "Verificado en el repo. Ya estaba implementado antes y no fue modificado en esta iteracion.")
    add_item(document, "1.3 ExpensesApplicationService - extendido + LocalStore.deleteExpense", True, "Verificado en el repo. Ya estaba implementado antes y no fue modificado en esta iteracion.")
    add_item(document, "1.4 ExpensesListView - con searchable + filtro categoria + swipeActions", True, "Verificado en el repo. Ya estaba implementado antes y no fue modificado en esta iteracion.")
    add_item(document, "1.5 ExpenseDetailView + EditExpenseView - separados (Detail read-only + Edit sheet)", True, "Verificado en el repo. Ya estaba implementado antes y no fue modificado en esta iteracion.")
    add_item(document, "Refactor TabView - Dashboard / Expenses / Insights", True, "Verificado en el repo. Se mantuvo tal como estaba; solo se cablearon dependencias nuevas de receipts hacia Expenses/Dashboard.")
    add_item(document, "1.6 EvC en F1 writes - URLError+EvC.swift helper + mensajes offline especificos en EditExpenseViewModel.save(), ExpenseDetailViewModel.delete(), ExpensesListViewModel.delete()", True, "Hecho en iOS. Implementado en Core/Support/URLError+EvC.swift, Features/Expenses/EditExpense/EditExpenseViewModel.swift, Features/Expenses/ExpenseDetail/ExpenseDetailViewModel.swift y Features/Expenses/ExpensesList/ExpensesListViewModel.swift.")

    document.add_heading("F2 - Photos + Caching - Companeros", level=2)
    add_item(document, "(Companeros) Bucket Supabase Storage receipts", True, "Hecho. Se agrego la migracion supabase/migrations/20260521_receipts_storage.sql con bucket privado y politicas por usuario autenticado.")
    add_item(document, "(Companeros) ReceiptsAdapter", True, "Hecho. Se agrego Core/Adapters/Supabase/ReceiptsAdapter.swift para upload/download autenticado de recibos.")
    add_item(document, "(Companeros) ImageCacheService (NSCache 50 MB / 200 items + disk cache 30 d TTL)", True, "Hecho. Se agrego Core/Support/ImageCacheService.swift con cache en memoria y disco, limite de 50 MB, 200 items y TTL de 30 dias.")
    add_item(document, "(Companeros) Integracion en LogExpense / ExpenseDetail", True, "Hecho. Se integro en Features/Expenses/LogExpense/* y Features/Expenses/ExpenseDetail/*, mas el cableado en AppContainer/MainTabView/ExpensesListView.")
    add_item(document, "(Companeros) EvC en photos (cache offline lectura)", True, "Hecho. ReceiptImageService prioriza cache local, mantiene pending uploads y permite fallback offline; LogExpenseViewModel y ExpenseDetailViewModel muestran mensajes adecuados.")

    document.add_heading("F3 - Insights (Fase 2 - Gabriel)", level=2)
    add_item(document, "2.1 Backend - 3 edge functions deployed ( get-bq-category-cycle-comparison , get-bq-category-streaks , get-bq-biggest-expense-of-cycle )", True, "Verificado en el repo. Ya estaba implementado antes y no se modifico.")
    add_item(document, "2.2 DTOs + FunctionsAdapter - InsightsModels.swift + 3 metodos en FunctionsAdapter", True, "Verificado en el repo. Ya estaba implementado antes y no se modifico.")
    add_item(document, "2.3 InsightSnapshot @Model - SwiftData con bqType , userId , payload , computedAt", True, "Verificado en el repo. Ya estaba implementado antes y no se modifico.")
    add_item(document, "2.4 InsightsApplicationService - async let paralelo + TTL 1h + EvC fallback", True, "Verificado en el repo. Ya estaba implementado antes y no se modifico.")
    add_item(document, "2.5 InsightsView + ViewModel + 3 Cards - wire-up completo en MainTabView", True, "Verificado en el repo. Ya estaba implementado antes y no se modifico.")
    add_item(document, "2.6 Polish - timestamp Updated X ago + forceRefresh + cleanup snapshots > 7 dias", True, "Verificado en el repo. Ya estaba implementado antes y no se modifico.")
    add_item(document, "Bug fix - race condition .task / .refreshable (verificar bundle == nil en .task )", True, "Verificado en el repo. Ya estaba implementado antes y no se modifico.")
    add_item(document, "Bug fix - URLError.cancelled con Task.detached en refresh()", True, "Verificado en el repo. Ya estaba implementado antes y no se modifico.")
    add_item(document, "Bug fix - catch separado URLError vs catch general (offline detection precisa)", True, "Verificado en el repo. Ya estaba implementado antes y no se modifico.")
    add_item(document, "Cleanup final - quitar prints [DIAG] , [VM/DEBUG] , [Service/DEBUG]", True, "Verificado en el repo. No se encontraron estos prints/tag de debug activos en iOS.")

    document.add_heading("Estrategias transversales (puntos 4.a, 4.b, 4.c, 4.d)", level=2)
    add_table(
        document,
        ["Punto", "Estrategia del PDF", "Feature", "Estado actualizado", "Comentario"],
        [
            ["4.a Multi-threading (20 pts)", "Swift Concurrency (async let) sobre GCD", "F3 InsightsApplicationService", "Hecho", "Sigue resuelto por F3 Insights."],
            ["4.b Local storage (20 pts)", "Database (SwiftData) + Expiration policy", "F3 InsightSnapshot", "Hecho", "Sigue resuelto por snapshots SwiftData; ademas hay drafts/pending receipts."],
            ["4.c Caching (20 pts)", "NSCache RAM + disk cache", "F2 ImageCacheService", "Hecho", "Quedo resuelto en iOS con cache RAM/disco para receipts."],
            ["4.d Eventual connectivity (20 pts)", "Cache, falling back to network + Expiration policy", "F3 Insights + F1 writes + F2 photos", "Hecho", "F3 ya estaba; F1.6 y F2 quedaron completos en esta iteracion."],
        ],
        [1.35, 1.85, 1.35, 1.0, 1.95],
    )

    document.add_heading("Punto 5 - Three Views (15 pts)", level=1)
    add_item(document, "View 1 - ExpensesListView", True, "Verificado en el repo. Ya estaba implementado antes y no se modifico estructuralmente.")
    add_item(document, "View 2 - ExpenseDetailView (con EditExpenseView como sheet)", True, "Verificado en el repo. Ya estaba implementado; en esta iteracion solo se amplio con receipts.")
    add_item(document, "View 3 - InsightsView (con 3 sub-cards)", True, "Verificado en el repo. Ya estaba implementado antes y no se modifico estructuralmente.")

    document.add_heading("Punto 6 - Three Business Questions (20 pts)", level=1)
    add_item(document, "BQ D - Category cycle comparison (current vs previous cycle)", True, "Verificado en el repo.")
    add_item(document, "BQ E - Top 3 category streaks (days without spending, capped at 30)", True, "Verificado en el repo.")
    add_item(document, "BQ F - Biggest expense of cycle + % of budget", True, "Verificado en el repo.")
    add_item(document, "Bonus - BQ C del Sprint 3 reparada (era global, ahora personal - coherente con D, E, F)", True, "Verificado en el repo.")

    document.add_heading("Punto 7 - Wiki Documentation", level=1)
    p = document.add_paragraph()
    set_para_spacing(p, before=0, after=6, line=1.1)
    run = p.add_run("Secciones obligatorias")
    set_run_font(run, 12, bold=True)
    add_item(document, "Features por sprint (S1, S2, S3, S4)", False, "Pendiente manual de wiki.")
    add_item(document, "Business Questions por sprint (A, B, C, D, E, F)", False, "Pendiente manual de wiki.")
    add_item(document, "Eventual Connectivity strategy - citar nombre exacto del PDF", False, "Pendiente manual de wiki. El contenido tecnico ya esta listo en InsightsApplicationService, URLError+EvC y ReceiptImageService.")
    add_bullet(document, 'Estrategia primaria: "Cache, falling back to network"')
    add_bullet(document, 'Estrategia secundaria: "Expiration policy" (TTL = 1h)')
    add_bullet(document, 'Strategies adicionales: "Pull / On user demand", "Cached on network response", "Generic fallback"')
    add_item(document, "Local Storage strategy - citar nombre exacto del PDF", False, "Pendiente manual de wiki. El contenido tecnico ya esta listo con SwiftData para insights y drafts/cache para receipts.")
    add_bullet(document, "Mecanismo: Database (SwiftData)")
    add_bullet(document, "Patron: Cache con Expiration policy")
    add_bullet(document, "Diferencia con Sprint 3: entity-cache -> computed-results-cache")
    add_item(document, "Multi-threading strategy - citar nombre exacto del PDF", False, "Pendiente manual de wiki. El contenido tecnico ya esta listo en InsightsApplicationService.")
    add_bullet(document, "GCD via Swift Concurrency (async let)")
    add_bullet(document, "Concurrent execution de 3 BQs paralelas")
    add_item(document, "Caching strategy - (companeros llenan)", False, "Pendiente manual de wiki. El contenido tecnico ya esta listo con ImageCacheService para F2.")
    add_item(document, "Value proposition refined", False, "Pendiente manual de wiki.")
    add_item(document, "Architecture overview (MVVM strict: View -> ViewModel -> ApplicationService -> Adapter -> Supabase)", False, "Pendiente manual de wiki.")
    add_item(document, "Diagrama de capas (opcional pero recomendado)", False, "Pendiente manual de wiki.")

    document.add_heading("Punto 8 - Ethics / Privacy Reflection", level=1)
    add_item(document, "Video Etica", False, "Pendiente manual. No existe evidencia de este entregable dentro del repo.")

    document.add_heading("Punto 9 - Firebase App Distribution", level=1)
    add_item(document, "Generar archive iOS (Product -> Archive en Xcode)", False, "Pendiente manual de release.")
    add_item(document, "Exportar IPA para distribucion ad-hoc", False, "Pendiente manual de release.")
    add_item(document, "Subir a Firebase App Distribution", False, "Pendiente manual de release.")
    add_item(document, "Crear tag git: isis3510-36-iOS-Sprint4", False, "Pendiente manual/release.")
    add_item(document, "Push del tag al repositorio", False, "Pendiente manual/release.")
    add_item(document, "Documentar link de descarga (Firebase) en Wiki", False, "Pendiente manual de release/wiki.")

    document.add_heading("Resumen visual del estado - actualizado", level=1)
    add_table(
        document,
        ["Seccion", "Estado actualizado", "Comentario"],
        [
            ["Backend + Edge functions", "100%", "Verificado; sin cambios en esta iteracion."],
            ["F1 CRUD", "100%", "F1.6 de eventual connectivity ya quedo completo."],
            ["F2 Photos + Caching", "100%", "Bucket, adapter, cache, integracion y EvC listos en iOS."],
            ["F3 Insights", "100%", "Verificado; cleanup final tambien confirmado."],
            ["Views (3)", "100%", "Las 3 vistas siguen completas; ExpenseDetail se amplio con receipts."],
            ["BQs (3 + bonus C)", "100%", "Verificado en el repo."],
            ["Estrategias 4.a, 4.b, 4.c, 4.d", "100% en codigo", "4.c y 4.d quedaron completas con F2 y F1.6."],
            ["Micro-optimization", "Codigo hecho", "Falta profiling ANTES/DESPUES y su documentacion."],
            ["Wiki", "Pendiente manual", "No puede cerrarse solo desde el repo."],
            ["Ethics reflection", "Pendiente manual", "No puede cerrarse solo desde el repo."],
            ["Firebase App Distribution", "Pendiente manual", "No puede cerrarse solo desde el repo."],
            ["Device real testing", "Pendiente manual", "Debe validarse en dispositivo fisico."],
        ],
        [2.2, 1.2, 3.1],
    )

    p = document.add_paragraph()
    set_para_spacing(p, before=4, after=0, line=1.0)
    run = p.add_run("Fin del checklist - Sprint 4 Fluxo G36")
    set_run_font(run, 11, italic=True)


def main():
    document = Document()
    build_styles(document)
    add_title(document)
    add_sections(document)
    document.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
